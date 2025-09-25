import streamlit as st
import pandas as pd
import numpy as np
import cv2
from ultralytics import YOLO
import tempfile
import os
from PIL import Image
import io
from datetime import datetime

# Imports for advanced reporting features
from docx import Document
from docx.shared import Inches
import plotly.express as px
from streamlit_image_comparison import image_comparison
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Image as PlatypusImage, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors

# --- Page config ---
st.set_page_config(page_title=" Marine-Foul-Detect", layout="wide")

# --- Paths ---
MODEL_PATH = "/Users/rudra/SY/CodeVerse/p2/runs/detect/yolov8_biofouling_fast3/weights/best.pt" 
LOGO_PATH = "Swanay.png" 

# --- Load model once ---
@st.cache_resource
def load_yolo_model(path):
    """Loads and caches the YOLO model."""
    try:
        return YOLO(path)
    except Exception as e:
        st.error(f"Failed to load YOLO model: {e}")
        return None

# --- Advanced Preprocessing ---
def preprocess_frame(frame, use_denoise, use_clahe, use_sharpen):
    """Applies selected enhancements to an image."""
    frame = frame.astype(np.uint8)
    processed = frame.copy()
    if use_denoise:
        processed = cv2.fastNlMeansDenoisingColored(processed, None, 10, 10, 7, 21)
    if use_clahe:
        lab = cv2.cvtColor(processed, cv2.COLOR_BGR2LAB)
        l, a, b = cv2.split(lab)
        clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8))
        cl = clahe.apply(l)
        merged = cv2.merge((cl, a, b))
        processed = cv2.cvtColor(merged, cv2.COLOR_LAB2BGR)
    if use_sharpen:
        kernel = np.array([[0, -1, 0], [-1, 5, -1], [0, -1, 0]])
        processed = cv2.filter2D(processed, -1, kernel)
    return processed.astype(np.uint8)

# --- Report Generation Functions ---
def create_word_report(orig_img, annotated_img, summary_df, details_df, coverage, severity, recommendation):
    """Creates a comprehensive Word document report."""
    document = Document()
    document.add_heading('Marine Foul Detection Report', level=1)
    document.add_paragraph(f"Report generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    # ... (rest of the Word function is the same as before)
    return document # (This function would be fully implemented as in previous versions)

def create_pdf_report(orig_img, annotated_img, summary_df, coverage, severity, recommendation):
    """Creates a comprehensive PDF report in an in-memory buffer."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []

    # Title
    story.append(Paragraph("Marine Fouling Detection Report", styles['h1']))
    story.append(Paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
    
    # Severity Analysis
    story.append(Paragraph("Severity Analysis", styles['h2']))
    story.append(Paragraph(f"<b>Overall Fouling Coverage:</b> {coverage:.2f}%", styles['Normal']))
    story.append(Paragraph(f"<b>Severity Level:</b> {severity}", styles['Normal']))
    story.append(Paragraph(f"<b>Recommendation:</b> {recommendation}", styles['Normal']))

    # Images
    story.append(Paragraph("Image Results", styles['h2']))
    orig_buffer = io.BytesIO()
    annotated_buffer = io.BytesIO()
    orig_img.save(orig_buffer, format="PNG")
    Image.fromarray(annotated_img).save(annotated_buffer, format="PNG")
    orig_buffer.seek(0)
    annotated_buffer.seek(0)
    
    im1 = PlatypusImage(orig_buffer, width=200, height=150)
    im2 = PlatypusImage(annotated_buffer, width=200, height=150)
    
    img_table = Table([[im1, im2]], colWidths=[220, 220])
    story.append(img_table)

    # Summary Table
    story.append(Paragraph("Detection Summary by Class", styles['h2']))
    table_data = [['Object Type', 'Count']] + summary_df.values.tolist()
    summary_table = Table(table_data)
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    story.append(summary_table)

    doc.build(story)
    buffer.seek(0)
    return buffer

# --- UI Layout ---
col_logo, col_title = st.columns([0.1, 0.9])
if os.path.exists(LOGO_PATH):
    with col_logo:
        st.image(LOGO_PATH, width=120)
with col_title:
    st.title("Marine-Foul-Detect")
    st.write("An advanced tool for detecting and quantifying marine fouling on ship hulls using YOLOv8.")

# --- Sidebar Controls ---
st.sidebar.header("⚙️ Analysis Controls")
uploaded_file = st.sidebar.file_uploader("Upload an underwater image", type=["jpg", "jpeg", "png"])
conf_threshold = st.sidebar.slider("Confidence Threshold", 0.0, 1.0, 0.25, 0.05)
with st.sidebar.expander("Advanced Preprocessing Options"):
    use_denoise = st.checkbox("Denoise", value=True)
    use_clahe = st.checkbox("Improve Contrast (CLAHE)", value=True)
    use_sharpen = st.checkbox("Sharpen", value=True)

# --- Main App Body ---
if uploaded_file is None:
    st.info("⬅️ Upload an image using the sidebar to start the analysis.")
else:
    # Use a temporary file for robust handling
    with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as tmp_in:
        tmp_in.write(uploaded_file.read())
        tmp_in_path = tmp_in.name

    orig_image = Image.open(tmp_in_path).convert("RGB")
    
    if st.button("🚀 Analyze Image", use_container_width=True):
        model = load_yolo_model(MODEL_PATH)
        if model is None:
            st.error("Model could not be loaded. Halting analysis.")
            st.stop()

        with st.spinner("Processing..."):
            frame = cv2.cvtColor(np.array(orig_image), cv2.COLOR_RGB2BGR)
            enhanced = preprocess_frame(frame, use_denoise, use_clahe, use_sharpen)
            results = model(enhanced)
            res = results[0]
            
            # Filter results and perform calculations
            filtered_boxes = [box for box in res.boxes if box.conf[0] >= conf_threshold]
            fouling_area = 0
            total_area = orig_image.width * orig_image.height

            rows = []
            for box in filtered_boxes:
                cls_id = int(box.cls[0])
                conf = float(box.conf[0])
                name = model.names[cls_id]
                rows.append([name, f"{conf:.2f}"])
                x1, y1, x2, y2 = box.xyxy[0]
                fouling_area += (x2 - x1) * (y2 - y1)
            
            details_df = pd.DataFrame(rows, columns=["Object Type", "Confidence"])
            
            # Calculate severity
            coverage = (fouling_area / total_area) * 100 if total_area > 0 else 0
            if coverage < 10:
                severity = "Light"
                recommendation = "Monitor condition. Schedule next cleaning in 2–3 months to maintain performance."
            elif coverage < 30:
                severity = "Moderate"
                recommendation = "Cleaning recommended to prevent significant fuel efficiency loss and hull damage."
            else:
                severity = "Heavy"
                recommendation = "Immediate cleaning is critical. Hull performance is significantly degraded."
            
            annotated = res.plot()
            annotated_rgb = cv2.cvtColor(annotated.astype(np.uint8), cv2.COLOR_BGR2RGB)
            
            # --- Display Results in Tabs ---
            tab1, tab2, tab3 = st.tabs(["📊 Summary & Report", "🖼️ Image Comparison", "📋 Detailed Detections"])

            with tab1:
                st.subheader("Analysis Summary")
                col1, col2, col3 = st.columns(3)
                col1.metric("Total Detections", len(details_df))
                col2.metric("Fouling Coverage", f"{coverage:.2f}%")
                col3.metric("Severity Level", severity)
                st.info(f"**Recommendation:** {recommendation}")
                st.markdown("---")

                if not details_df.empty:
                    summary_counts = details_df['Object Type'].value_counts().reset_index()
                    summary_counts.columns = ['Object Type', 'Count']
                    
                    fig = px.bar(summary_counts, x='Object Type', y='Count', title='Detection Counts per Class')
                    st.plotly_chart(fig, use_container_width=True)
                    
                    # Report Downloads
                    st.subheader("Download Reports")
                    col_pdf, col_csv = st.columns(2)
                    
                    with col_pdf:
                        pdf_buffer = create_pdf_report(orig_image, annotated_rgb, summary_counts, coverage, severity, recommendation)
                        st.download_button(
                            label="📥 Download PDF Report",
                            data=pdf_buffer,
                            file_name=f"fouling_report_{datetime.now().strftime('%Y%m%d')}.pdf",
                            mime="application/pdf"
                        )
                    
                    with col_csv:
                        csv_buffer = details_df.to_csv(index=False).encode('utf-8')
                        st.download_button(
                            "📥 Download Detailed Data (CSV)",
                            data=csv_buffer,
                            file_name=f"detailed_detections_{datetime.now().strftime('%Y%m%d')}.csv",
                            mime="text/csv"
                        )
                else:
                    st.success("✅ No fouling detected above the confidence threshold. Hull is clean!")

            with tab2:
                st.subheader("Visual Comparison")
                image_comparison(
                    img1=orig_image, label1="Original",
                    img2=annotated_rgb, label2="Annotated"
                )

            with tab3:
                st.subheader("Detailed Detection Data")
                st.dataframe(details_df, use_container_width=True)

    # Clean up the temporary file
    if os.path.exists(tmp_in_path):
        os.remove(tmp_in_path)